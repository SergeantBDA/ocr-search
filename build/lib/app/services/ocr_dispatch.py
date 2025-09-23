from typing import Optional, Tuple
from io import BytesIO
from pathlib import Path
import os

from .ocr_pdf import extract_text_from_pdf
from .ocr_image import extract_text_from_image
from .metadata import extract_metadata  # new import


def extract_text(filename: str, content: bytes, mime: Optional[str]) -> Tuple[str, dict]:
    """
    Выбирает обработчик по MIME или расширению и возвращает (text, meta).
    Поддерживаемые форматы: pdf, jpg/jpeg, png, docx, xlsx.
    При несоответствии формата поднимает ValueError.
    """
    mime = (mime or "").lower()
    ext = Path(filename).suffix.lower()

    # PDF
    if mime == "application/pdf" or ext == ".pdf":
        text = extract_text_from_pdf(content)
        meta = extract_metadata(filename, content)
        return text, meta

    # Images
    if mime.startswith("image/") or ext in {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}:
        text = extract_text_from_image(content)
        meta = extract_metadata(filename, content)
        return text, meta

    # DOCX
    if mime in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    } or ext in {".docx", ".doc"}:
        try:
            from docx import Document as DocxDocument
        except Exception as e:
            raise ValueError("python-docx is required to process docx files") from e

        try:
            bio = BytesIO(content)
            doc = DocxDocument(bio)
            parts = []

            # paragraphs
            for p in doc.paragraphs:
                text = (p.text or "").strip()
                if text:
                    parts.append(text)

            # tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join((cell.text or "").strip() for cell in row.cells if (cell.text or "").strip())
                    if row_text:
                        parts.append(row_text)

            full_text = "\n\n".join(parts).strip()
            meta = extract_metadata(filename, content)
            return full_text, meta
        except Exception:
            return "", {}

    # XLSX
    if mime in {
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    } or ext in {".xlsx", ".xls"}:
        try:
            import openpyxl
        except Exception as e:
            raise ValueError("openpyxl is required to process xlsx files") from e

        try:
            bio = BytesIO(content)
            wb = openpyxl.load_workbook(filename=bio, data_only=True, read_only=True)
            parts = []
            for sheet in wb.worksheets:
                sheet_parts = []
                for row in sheet.iter_rows(values_only=True):
                    row_vals = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                    if row_vals:
                        sheet_parts.append(" ".join(row_vals))
                if sheet_parts:
                    parts.append("\n".join(sheet_parts))
            full_text = "\n\n".join(parts).strip()
            meta = extract_metadata(filename, content)
            return full_text, meta
        except Exception:
            return "", {}

    #raise ValueError(f"Unsupported file type (mime={mime!r}, ext={ext!r})")
