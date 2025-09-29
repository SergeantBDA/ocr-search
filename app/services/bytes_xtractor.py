"""
Извлечение текста из файлов, переданных как bytes.
Работает по mime-типу (приоритет) или по расширению из filename.
Спроектирован для параллельного запуска (ThreadPoolExecutor friendly).

Зависимости (опциональные части включены в try/except):
- PyMuPDF (fitz)                 — PDF, рендер для OCR fallback
- pytesseract + Tesseract OCR    — OCR (изображения, PDF fallback)
- Pillow (PIL)                   — работа с изображениями
- python-docx                    — docx
- pandas + openpyxl/xlrd         — excel
- beautifulsoup4 + lxml          — html
- striprtf                       — rtf
"""

from __future__ import annotations

import io
import logging
import mimetypes
from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import Optional, Dict, Type

# --------- optional imports guarded for thread-safe lazy usage ----------
try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

try:
    from PIL import Image
except Exception:  # pragma: no cover
    Image = None

try:
    import pytesseract
except Exception:  # pragma: no cover
    pytesseract = None

try:
    from bs4 import BeautifulSoup
except Exception:  # pragma: no cover
    BeautifulSoup = None

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None

try:
    import pandas as pd
except Exception:  # pragma: no cover
    pd = None

try:
    from striprtf.striprtf import rtf_to_text
except Exception:  # pragma: no cover
    rtf_to_text = None

from email import policy
from email.parser import BytesParser

# ------------------------- logging --------------------------------------
from app.logger import logger as app_logger, attach_to_logger_names
attach_to_logger_names(["app.services.bytes_xtractor"])

# ----------------------- helpers / heuristics ---------------------------
RUSSIAN_CHARS = set(r".:,-+=()!0123456789абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ")

def looks_like_russian(text: str, threshold: float = 0.40) -> bool:
    if not text:
        return False
    ru = sum(1 for ch in text if ch in RUSSIAN_CHARS)
    return (ru / max(1, len(text))) >= threshold

def normalized_mime(mime: Optional[str]) -> Optional[str]:
    return mime.lower().strip() if mime else None

def ext_from_filename(filename: Optional[str]) -> Optional[str]:
    if not filename:
        return None
    dot = filename.rfind(".")
    if dot == -1:
        return None
    return filename[dot + 1 :].lower()

# --------------------------- base ---------------------------------------
@dataclass(frozen=True)
class BytesPayload:
    content: bytes
    filename: Optional[str] = None
    mime: Optional[str] = None

class BytesExtractor(ABC):
    """Базовый класс для экстракторов, работающих с bytes."""

    def __init__(self, payload: BytesPayload, *, ocr_lang: str = "rus+eng"):
        self.payload = payload
        self.ocr_lang = ocr_lang

    @abstractmethod
    def extract_text(self) -> str: ...

# -------------------------- concrete extractors -------------------------

class PDFExtractor(BytesExtractor):
    """PDF: сперва текст слоем, затем (если пусто/мало) OCR по страницам."""
    def extract_text(self) -> str:
        if fitz is None:
            app_logger.warning("PyMuPDF не установлен — пропускаю PDF.")
            return ""
        txt_parts = []
        doc = fitz.open(stream=self.payload.content, filetype="pdf")
        # 1) прямой текст
        for page in doc:
            t = page.get_text().strip()
            if t:
                txt_parts.append(t)
        whole_text = "\n".join(txt_parts).strip()
        #app_logger.info(f"PDF TEXT LEN:{len(whole_text)}")
        # 2) fallback OCR, если пусто или явно не русский при ожидании рус/eng контента
        if not whole_text or (not looks_like_russian(whole_text[:200]) and pytesseract and Image):
            ocr_parts = []
            for page in doc:
                pix = page.get_pixmap(dpi=300, alpha=False)
                img_bytes = pix.tobytes("png")
                if Image is None or pytesseract is None:
                    continue
                img = Image.open(io.BytesIO(img_bytes))
                ocr_parts.append(pytesseract.image_to_string(img, lang=self.ocr_lang))
            ocr_text = "\n".join(ocr_parts).strip()
            if len(ocr_text) > len(whole_text):
                return ocr_text
        return whole_text

class ImageExtractor(BytesExtractor):
    def extract_text(self) -> str:
        if Image is None or pytesseract is None:
            app_logger.warning("Pillow/pytesseract не установлены — пропускаю изображение.")
            return ""
        img = Image.open(io.BytesIO(self.payload.content))
        return pytesseract.image_to_string(img, lang=self.ocr_lang)

class EMLExtractor(BytesExtractor):
    def extract_text(self) -> str:
        try:
            msg = BytesParser(policy=policy.default).parsebytes(self.payload.content)
        except Exception as e:
            app_logger.exception("Ошибка парсинга EML: %s", e)
            return ""
        lines = []
        def _safe(val): return "" if val is None else str(val)
        lines.append(f"Тема: {_safe(msg['subject'])}")
        lines.append(f"От  : {_safe(msg['from'])}")
        lines.append(f"Кому: {_safe(msg['to'])}")
        lines.append(f"Дата: {_safe(msg['date'])}")

        text_part = ""
        html_part = ""
        for part in msg.walk():
            ctype = part.get_content_type()
            dispo = part.get_content_disposition()
            if dispo == "attachment":
                continue
            if ctype == "text/plain":
                try:
                    text_part = part.get_content()
                except Exception:
                    pass
            elif ctype == "text/html":
                try:
                    html_part = part.get_content()
                except Exception:
                    pass

        body = text_part or html_part
        if body:
            lines.append("Тело письма:\n" + body)

        # перечислим вложения
        attachments = [att.get_filename() or "attachment" for att in msg.iter_attachments()]
        if attachments:
            lines.append("Вложения:\n" + "\n".join(f"{i+1}. {n}" for i, n in enumerate(attachments)))
        return "\n".join(lines)

class RTFExtractor(BytesExtractor):
    def extract_text(self) -> str:
        if rtf_to_text is None:
            app_logger.warning("striprtf не установлен — пропускаю RTF.")
            return ""
        try:
            rtf_content = self.payload.content.decode("utf-8", errors="ignore")
        except Exception:
            rtf_content = ""
        return rtf_to_text(rtf_content) if rtf_content else ""

class HTMLExtractor(BytesExtractor):
    def extract_text(self) -> str:
        if BeautifulSoup is None:
            app_logger.warning("beautifulsoup4 не установлен — пропускаю HTML.")
            return ""
        html = self._decode_text_like()
        soup = BeautifulSoup(html, "lxml") if html else None
        if not soup:
            return ""
        for tag in soup(["script","style","nav","header","footer","aside","noscript","link","meta","form"]):
            tag.decompose()
        return soup.get_text(separator=" ", strip=True)

    def _decode_text_like(self) -> str:
        # подбираем распространённые кодировки
        for enc in ("utf-8","cp1251","koi8-r","utf-16","iso-8859-5","mac-cyrillic"):
            try:
                return self.payload.content.decode(enc)
            except UnicodeDecodeError:
                continue
        return self.payload.content.decode("utf-8", errors="ignore")

class XMLExtractor(HTMLExtractor):
    """Для простоты используем ту же стратегию что и для HTML (strip-текст)."""
    pass

class DOCXExtractor(BytesExtractor):
    def extract_text(self) -> str:
        if DocxDocument is None:
            app_logger.warning("python-docx не установлен — пропускаю DOCX.")
            return ""
        bio = io.BytesIO(self.payload.content)
        try:
            doc = DocxDocument(bio)
        except Exception as e:
            app_logger.exception("Ошибка чтения docx: %s", e)
            return ""
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()

class ExcelExtractor(BytesExtractor):
    def extract_text(self) -> str:
        if pd is None:
            app_logger.warning("pandas не установлен — пропускаю Excel.")
            return ""
        bio = io.BytesIO(self.payload.content)
        try:
            # sheet_name=None => dict листов
            frames: Dict[str, "pd.DataFrame"] = pd.read_excel(bio, sheet_name=None, header=None)
        except Exception as e:
            app_logger.exception("Ошибка чтения Excel: %s", e)
            return ""
        parts = []
        for sheet, df in frames.items():
            parts.append(f"=== Лист: {sheet} ===")
            parts.append(df.to_csv(sep="\t", index=False, header=False))
        return "\n\n".join(parts).strip()

class PlainTextExtractor(BytesExtractor):
    def extract_text(self) -> str:
        for enc in ("utf-8","cp1251","koi8-r","utf-16","iso-8859-5","mac-cyrillic"):
            try:
                return self.payload.content.decode(enc)
            except UnicodeDecodeError:
                continue
        return self.payload.content.decode("utf-8", errors="ignore")

class UnsupportedExtractor(BytesExtractor):
    def extract_text(self) -> str:
        app_logger.info("Неподдерживаемый тип: %s / %s",
                    self.payload.mime, self.payload.filename)
        return ""

# --------------------------- factory ------------------------------------

_MIME_MAP: Dict[str, Type[BytesExtractor]] = {
    "application/pdf": PDFExtractor,
    "image/png": ImageExtractor,
    "image/jpeg": ImageExtractor,
    "image/bmp": ImageExtractor,
    "image/gif": ImageExtractor,
    "text/html": HTMLExtractor,
    "application/xhtml+xml": HTMLExtractor,
    "text/xml": XMLExtractor,
    "application/xml": XMLExtractor,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DOCXExtractor,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ExcelExtractor,
    "text/plain": PlainTextExtractor,
    "message/rfc822": EMLExtractor,
}

_EXT_MAP: Dict[str, Type[BytesExtractor]] = {
    "pdf": PDFExtractor,
    "png": ImageExtractor,
    "jpg": ImageExtractor,
    "jpeg": ImageExtractor,
    "bmp": ImageExtractor,
    "gif": ImageExtractor,
    "html": HTMLExtractor,
    "htm": HTMLExtractor,
    "xml": XMLExtractor,
    "docx": DOCXExtractor,
    "xlsx": ExcelExtractor,
    "xls": ExcelExtractor,
    "xlsm": ExcelExtractor,
    "rtf": RTFExtractor,
    "txt": PlainTextExtractor,
    "eml": EMLExtractor,
    # .msg — как правило требует extract_msg с файловым путем; пропустим в bytes-варианте
}

def _guess_mime_from_name(filename: Optional[str]) -> Optional[str]:
    if not filename:
        return None
    guessed, _ = mimetypes.guess_type(filename)
    return guessed

def get_extractor(payload: BytesPayload, *, ocr_lang: str = "rus+eng") -> BytesExtractor:
    """
    Выбор экстрактора:
    1) mime из payload.mime (если указан)
    2) mime по mimetypes.guess_type(filename)
    3) расширение из filename
    4) fallback: PlainText/Unsupported
    """
    mime = normalized_mime(payload.mime) or normalized_mime(_guess_mime_from_name(payload.filename))
    if mime and mime in _MIME_MAP:
        return _MIME_MAP[mime](payload, ocr_lang=ocr_lang)

    ext = ext_from_filename(payload.filename)
    if ext and ext in _EXT_MAP:
        return _EXT_MAP[ext](payload, ocr_lang=ocr_lang)

    # Если неизвестно, попробуем как текст:
    return UnsupportedExtractor(payload, ocr_lang=ocr_lang)

# --------------------------- public API ----------------------------------

def extract_text_bytes(
    content: bytes,
    filename: Optional[str] = None,
    mime: Optional[str] = None,
    *,
    ocr_lang: str = "rus+eng",
) -> str:
    """
    Удобная функция «в один вызов». Потокобезопасна.
    Пример использования с ThreadPoolExecutor:

        tasks = []
        with ThreadPoolExecutor() as tp:
            for item in items:
                tasks.append(tp.submit(extract_text_bytes, item.content, item.name, item.mime))
        results = [t.result() for t in tasks]

    """
    payload = BytesPayload(content=content, filename=filename, mime=mime)
    extractor = get_extractor(payload, ocr_lang=ocr_lang)
    return extractor.extract_text()
