# bytes_xtractor.py
"""
Извлечение текста из файлов, переданных как bytes.
Учитывает ориентацию страниц при OCR (pytesseract.image_to_osd).

Приоритет стратегии: MIME (если передан) → mimetypes по filename → расширение → plain text.

Опциональные зависимости:
- PyMuPDF (fitz)                 — PDF (текстовый слой + OCR-фоллбек)
- pytesseract + Tesseract OCR    — OCR (изображения, PDF-фоллбек)
- Pillow (PIL)                   — изображения для OCR
- python-docx                    — DOCX
- pandas + openpyxl/xlrd         — Excel
- beautifulsoup4 + lxml          — HTML/XML
- striprtf                       — RTF
"""

from __future__ import annotations
import re
import io
import logging
import mimetypes
from abc import ABC, abstractmethod
from dataclasses import dataclass
from typing import Optional, Dict, Type

# -------- optional imports (защищены) ----------
try:
    import fitz  # PyMuPDF
except Exception:  # pragma: no cover
    fitz = None

try:
    from PIL import Image
except Exception:  # pragma: no cover
    Image = None

try:
    import pytesseract
except Exception:  # pragma: no cover
    pytesseract = None

try:
    from bs4 import BeautifulSoup
except Exception:  # pragma: no cover
    BeautifulSoup = None

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None

try:
    import pandas as pd
except Exception:  # pragma: no cover
    pd = None

try:
    from striprtf.striprtf import rtf_to_text
except Exception:  # pragma: no cover
    rtf_to_text = None

from email import policy
from email.parser import BytesParser

# ------------------------- logging --------------------------------------
from app.logger import logger as app_logger, attach_to_logger_names
attach_to_logger_names(["app.services.bytes_xtractor"])
# ----------------------- helpers ----------------------------------------
RUSSIAN_CHARS = set(r".:,-+=()!0123456789абвгдеёжзийклмнопрстуфхцчшщъыьэюяАБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ")

def looks_like_russian(text: str, threshold: float = 0.40) -> bool:
    if not text:
        return False
    ru = sum(1 for ch in text if ch in RUSSIAN_CHARS)
    return (ru / max(1, len(text))) >= threshold

def normalized_mime(mime: Optional[str]) -> Optional[str]:
    return mime.lower().strip() if mime else None

def ext_from_filename(filename: Optional[str]) -> Optional[str]:
    if not filename or "." not in filename:
        return None
    return filename.rsplit(".", 1)[-1].lower()

def _guess_mime_from_name(filename: Optional[str]) -> Optional[str]:
    if not filename:
        return None
    guessed, _ = mimetypes.guess_type(filename)
    return guessed

def _rotate_by_osd(
    img: "Image.Image",
    *,
    enable_trials: bool = False,           # можно включить дешёвую эвристику «перебор углов», если OSD упал
    trial_angles: tuple[int, ...] = (0, 90, 180, 270),
    osd_conf_threshold: float = 3.0,       # ниже — считаем, что OSD не уверен
    logger: logging.Logger = app_logger,
) -> tuple["Image.Image", dict]:
    """
    Пытается определить ориентацию через Tesseract OSD и повернуть изображение.
    Возвращает (возможно повернутое) изображение и метаданные:
      {"method": "osd"|"trial"|"skip", "angle": int, "confidence": float, "reason": str}

    НИКОГДА тихо не молчит: всегда пишет в лог на уровне INFO/WARNING/ERROR.
    """
    meta = {"method": "skip", "angle": 0, "confidence": -1.0, "reason": ""}

    if pytesseract is None:
        msg = "pytesseract не установлен — пропускаю OSD."
        logger.info(msg)
        meta["reason"] = msg
        return img, meta

    # Pillow: нормализуем режим — OSD иногда капризничает на 'P', 'LA', 'RGBA'
    try:
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
    except Exception as e:
        logger.warning("Не удалось сконвертировать изображение к RGB для OSD: %s", e)

    w, h = getattr(img, "size", (0, 0))
    if min(w, h) < 50:
        msg = f"слишком маленькое изображение для OSD: {w}x{h}"
        logger.info(msg)
        meta["reason"] = msg
        return img, meta

    # Явная конфигурация OSD
    try:
        osd_text = pytesseract.image_to_osd(img, config="--psm 0 -l osd")
        # Пример:
        # Orientation in degrees: 270
        # Rotate: 90
        # Orientation confidence: 12.34
        # Script: Cyrillic
        # Script confidence: 9.87

        # Парсим angle
        angle = None
        m = re.search(r"(?i)Rotate:\s*(\d+)", osd_text)
        if m:
            angle = int(m.group(1))
        else:
            m2 = re.search(r"(?i)Orientation in degrees:\s*(\d+)", osd_text)
            if m2:
                # В некоторых сборках «Orientation in degrees» = фактический угол страницы.
                # Для Pillow поворачиваем на -angle (против часовой, чтобы выровнять).
                angle = (-int(m2.group(1))) % 360

        # Парсим confidence
        conf = -1.0
        c = re.search(r"(?i)Orientation confidence:\s*([0-9]+(?:\.[0-9]+)?)", osd_text)
        if c:
            conf = float(c.group(1))

        logger.info("OSD результат: angle=%s, confidence=%.2f (size=%dx%d)", angle, conf, w, h)

        if angle is None:
            meta.update({"method": "osd", "angle": 0, "confidence": conf, "reason": "angle not parsed"})
            logger.info("OSD не смог распознать угол — пропускаю поворот.")
            return img, meta

        if conf < osd_conf_threshold:
            logger.info("OSD низкая уверенность (%.2f < %.2f)", conf, osd_conf_threshold)

        # Приводим угол к {0,90,180,270}
        angle = int(angle) % 360
        if angle in (0, 360):
            meta.update({"method": "osd", "angle": 0, "confidence": conf, "reason": "no rotation needed"})
            return img, meta

        # Pillow: положительный угол — поворот против часовой стрелки.
        # OSD «Rotate: 90» обычно означает «поверни по часовой на 90», значит надо повернуть на -90.
        # Мы уже нормализовали выше, здесь просто -angle:
        rotated = img.rotate(-angle, expand=True)
        meta.update({"method": "osd", "angle": angle, "confidence": conf, "reason": "rotated by OSD"})
        logger.info("Повернул по OSD на %d°", angle)
        return rotated, meta

    except Exception as e:
        # Логируем ПРИЧИНУ (например, TesseractNotFoundError, отсутствие osd.traineddata)
        logger.error("OSD упал: %s", e)
        meta["reason"] = f"osd failed: {e}"

    # === Fallback (опционально): перебор углов, если OSD не сработал ===
    if enable_trials:
        best_angle = 0
        best_score = -1
        best_img = img
        logger.info("Пробую эвристику углов: %s", trial_angles)

        for a in trial_angles:
            try:
                candidate = img if a in (0, 360) else img.rotate(-a, expand=True)
                txt = pytesseract.image_to_string(candidate, lang="rus+eng", config="--psm 6")
                # Простейший скорер «сколько букв/цифр»:
                letters = sum(ch.isalnum() for ch in txt)
                # Можно усилить: штрафовать за «слишком мало» или «слишком много» пробелов/непечатаемых
                score = letters
                logger.info("trial angle=%d → score=%d, len=%d", a, score, len(txt))
                if score > best_score:
                    best_score = score
                    best_angle = a
                    best_img = candidate
            except Exception as e:
                logger.warning("trial angle=%d упал: %s", a, e)

        meta.update({"method": "trial", "angle": best_angle, "confidence": -1.0, "reason": "trial best angle"})
        if best_angle not in (0, 360):
            logger.info("Выбрал эвристический угол %d° (score=%d)", best_angle, best_score)
        return best_img, meta

    return img, meta

# --------------------------- base ---------------------------------------
@dataclass(frozen=True)
class BytesPayload:
    content: bytes
    filename: Optional[str] = None
    mime: Optional[str] = None

class BytesExtractor(ABC):
    """Базовый класс для экстракторов, работающих с bytes."""

    def __init__(self, payload: BytesPayload, *, ocr_lang: str = "rus+eng"):
        self.payload = payload
        self.ocr_lang = ocr_lang

    @abstractmethod
    def extract_text(self) -> str: ...

# -------------------------- concrete extractors -------------------------

class PDFExtractor(BytesExtractor):
    """PDF: сперва текстовый слой, затем OCR-фоллбек с учётом ориентации."""
    def extract_text(self) -> str:
        if fitz is None:
            app_loger.warning("PyMuPDF не установлен — пропускаю PDF.")
            return ""
        try:
            doc = fitz.open(stream=self.payload.content, filetype="pdf")
        except Exception as e:
            app_loger.exception("Ошибка открытия PDF: %s", e)
            return ""

        # 1) прямой текст
        txt_parts = []
        for page in doc:
            try:
                t = page.get_text().strip()
            except Exception:
                t = ""
            if t:
                txt_parts.append(t)
        whole_text = "\n".join(txt_parts).strip()

        # 2) OCR fallback при отсутствии/скупости текста
        need_ocr = not whole_text or (not looks_like_russian(whole_text[:300]) and pytesseract and Image)
        if need_ocr and Image is not None and pytesseract is not None:
            ocr_parts = []
            for page in doc:
                try:                    
                    pix = page.get_pixmap(dpi=300, alpha=False)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    img = _rotate_by_osd(img)[0]                    
                    ocr_parts.append(pytesseract.image_to_string(img, lang=self.ocr_lang))
                except Exception:
                    # продолжаем собирать со следующих страниц
                    continue            
            ocr_text = "\n".join(ocr_parts).strip()
            if len(ocr_text) > len(whole_text):
                return ocr_text

        return whole_text

class ImageExtractor(BytesExtractor):
    """OCR изображений с определением ориентации (OSD)."""
    def extract_text(self) -> str:
        if Image is None or pytesseract is None:
            app_loger.warning("Pillow/pytesseract не установлены — пропускаю изображение.")
            return ""
        try:
            img = Image.open(io.BytesIO(self.payload.content))
        except Exception as e:
            app_loger.exception("Ошибка открытия изображения: %s", e)
            return ""
        img = _rotate_by_osd(img)
        try:
            return pytesseract.image_to_string(img, lang=self.ocr_lang)
        except Exception as e:
            app_loger.exception("Ошибка OCR изображения: %s", e)
            return ""

class EMLExtractor(BytesExtractor):
    def extract_text(self) -> str:
        try:
            msg = BytesParser(policy=policy.default).parsebytes(self.payload.content)
        except Exception as e:
            app_loger.exception("Ошибка парсинга EML: %s", e)
            return ""
        lines = []
        def _safe(val): return "" if val is None else str(val)
        lines.append(f"Тема: {_safe(msg['subject'])}")
        lines.append(f"От  : {_safe(msg['from'])}")
        lines.append(f"Кому: {_safe(msg['to'])}")
        lines.append(f"Дата: {_safe(msg['date'])}")

        text_part = ""
        html_part = ""
        for part in msg.walk():
            ctype = part.get_content_type()
            dispo = part.get_content_disposition()
            if dispo == "attachment":
                continue
            if ctype == "text/plain":
                try:
                    text_part = part.get_content()
                except Exception:
                    pass
            elif ctype == "text/html":
                try:
                    html_part = part.get_content()
                except Exception:
                    pass

        body = text_part or html_part
        if body:
            lines.append("Тело письма:\n" + body)

        attachments = [att.get_filename() or "attachment" for att in msg.iter_attachments()]
        if attachments:
            lines.append("Вложения:\n" + "\n".join(f"{i+1}. {n}" for i, n in enumerate(attachments)))
        return "\n".join(lines)

class RTFExtractor(BytesExtractor):
    def extract_text(self) -> str:
        if rtf_to_text is None:
            app_loger.warning("striprtf не установлен — пропускаю RTF.")
            return ""
        try:
            rtf_content = self.payload.content.decode("utf-8", errors="ignore")
            return rtf_to_text(rtf_content) if rtf_content else ""
        except Exception as e:
            app_loger.exception("Ошибка чтения RTF: %s", e)
            return ""

class HTMLExtractor(BytesExtractor):
    def extract_text(self) -> str:
        if BeautifulSoup is None:
            app_loger.warning("beautifulsoup4 не установлен — пропускаю HTML.")
            return ""
        html = self._decode_text_like()
        try:
            soup = BeautifulSoup(html, "lxml") if html else None
        except Exception:
            soup = None
        if not soup:
            return ""
        for tag in soup(["script","style","nav","header","footer","aside","noscript","link","meta","form"]):
            tag.decompose()
        return soup.get_text(separator=" ", strip=True)

    def _decode_text_like(self) -> str:
        for enc in ("utf-8","cp1251","koi8-r","utf-16","iso-8859-5","mac-cyrillic"):
            try:
                return self.payload.content.decode(enc)
            except UnicodeDecodeError:
                continue
        return self.payload.content.decode("utf-8", errors="ignore")

class XMLExtractor(HTMLExtractor):
    pass

class DOCXExtractor(BytesExtractor):
    def extract_text(self) -> str:
        if DocxDocument is None:
            app_loger.warning("python-docx не установлен — пропускаю DOCX.")
            return ""
        bio = io.BytesIO(self.payload.content)
        try:
            doc = DocxDocument(bio)
        except Exception as e:
            app_loger.exception("Ошибка чтения DOCX: %s", e)
            return ""
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()

class ExcelExtractor(BytesExtractor):
    def extract_text(self) -> str:
        if pd is None:
            app_loger.warning("pandas не установлен — пропускаю Excel.")
            return ""
        bio = io.BytesIO(self.payload.content)
        try:
            frames: Dict[str, "pd.DataFrame"] = pd.read_excel(bio, sheet_name=None, header=None)
        except Exception as e:
            app_loger.exception("Ошибка чтения Excel: %s", e)
            return ""
        parts = []
        for sheet, df in frames.items():
            parts.append(f"=== Лист: {sheet} ===")
            parts.append(df.to_csv(sep="\t", index=False, header=False))
        return "\n\n".join(parts).strip()

class PlainTextExtractor(BytesExtractor):
    def extract_text(self) -> str:
        for enc in ("utf-8","cp1251","koi8-r","utf-16","iso-8859-5","mac-cyrillic"):
            try:
                return self.payload.content.decode(enc)
            except UnicodeDecodeError:
                continue
        return self.payload.content.decode("utf-8", errors="ignore")

class UnsupportedExtractor(BytesExtractor):
    def extract_text(self) -> str:
        app_loger.info("Неподдерживаемый тип: %s / %s",
                    self.payload.mime, self.payload.filename)
        return ""

# --------------------------- factory ------------------------------------

_MIME_MAP: Dict[str, Type[BytesExtractor]] = {
    "application/pdf": PDFExtractor,
    "image/png": ImageExtractor,
    "image/jpeg": ImageExtractor,
    "image/bmp": ImageExtractor,
    "image/gif": ImageExtractor,
    "text/html": HTMLExtractor,
    "application/xhtml+xml": HTMLExtractor,
    "text/xml": XMLExtractor,
    "application/xml": XMLExtractor,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DOCXExtractor,
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ExcelExtractor,
    "text/plain": PlainTextExtractor,
    "message/rfc822": EMLExtractor,
}

_EXT_MAP: Dict[str, Type[BytesExtractor]] = {
    "pdf": PDFExtractor,
    "png": ImageExtractor,
    "jpg": ImageExtractor,
    "jpeg": ImageExtractor,
    "bmp": ImageExtractor,
    "gif": ImageExtractor,
    "html": HTMLExtractor,
    "htm": HTMLExtractor,
    "xml": XMLExtractor,
    "docx": DOCXExtractor,
    "xlsx": ExcelExtractor,
    "xls": ExcelExtractor,
    "xlsm": ExcelExtractor,
    "rtf": RTFExtractor,
    "txt": PlainTextExtractor,
    "eml": EMLExtractor,
}

def get_extractor(payload: BytesPayload, *, ocr_lang: str = "rus+eng") -> BytesExtractor:
    """
    Выбор экстрактора:
      1) MIME из payload.mime,
      2) MIME по mimetypes.guess_type(filename),
      3) расширение из filename,
      4) fallback: UnsupportedExtractor.
    """
    mime = normalized_mime(payload.mime) or normalized_mime(_guess_mime_from_name(payload.filename))
    if mime and mime in _MIME_MAP:
        return _MIME_MAP[mime](payload, ocr_lang=ocr_lang)

    ext = ext_from_filename(payload.filename)
    if ext and ext in _EXT_MAP:
        return _EXT_MAP[ext](payload, ocr_lang=ocr_lang)

    return UnsupportedExtractor(payload, ocr_lang=ocr_lang)

# --------------------------- public API ----------------------------------

def extract_text_bytes(
    content: bytes,
    filename: Optional[str] = None,
    mime: Optional[str] = None,
    *,
    ocr_lang: str = "rus+eng",
) -> str:
    """
    Удобная функция «в один вызов». Потокобезопасна.

    Пример с ThreadPoolExecutor:

        with ThreadPoolExecutor(max_workers=4) as tp:
            futs = [tp.submit(extract_text_bytes, b, name, mime) for (b,name,mime) in items]
            results = [f.result() for f in futs]
    """
    payload = BytesPayload(content=content, filename=filename, mime=mime)
    extractor = get_extractor(payload, ocr_lang=ocr_lang)
    return extractor.extract_text()
