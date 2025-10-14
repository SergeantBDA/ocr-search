from __future__ import annotations
import io
import logging
from pathlib import Path
from .base import BytesExtractor
# ------------------------- logging --------------------------------------
from app.logger_worker import worker_log as app_logger

try:
    from docx import Document as DocxDocument
except Exception:  # pragma: no cover
    DocxDocument = None

class DOCXExtractor(BytesExtractor):
    def extract_text(self) -> str:
        if DocxDocument is None:
            app_logger.warning("python-docx не установлен — пропускаю DOCX.")
            return ""
        
        try:
            if self.payload.path:
                doc = DocxDocument(self.payload.path)
            else:
                bio = io.BytesIO(self.payload.content or b"")
                doc = DocxDocument(bio)
        except Exception as e:
            app_logger.exception("Ошибка чтения DOCX: %s", e)
            return ""
        parts = []
        for p in doc.paragraphs:
            parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts).strip()